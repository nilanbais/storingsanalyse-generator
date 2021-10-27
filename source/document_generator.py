"""
Class voor het genereren van de tekst voor de verschillende projecten.

Hoe dit precies wordt opgebouwd is work-in-progress -> het kan dat er per project een specifieke class geschreven
moet worden, maar dat deze gebruik maken van één overkoepelende class die de meest generalistische functies bevat.

Voor de Coentunnel wordt begonnen met het opbouwen van een class die specifiek voor de Coentunnel is. In de loop
van dit process wordt er gekeken naar de mogelijkheden om het eerder beschreven idee te realiseren
"""
from storingsanalyse import StoringsAnalyse
from pandas import DataFrame
import numpy as np
import docx
import os
from typing import Tuple, Iterable, List

# todo: onderstaande verwerken
"""
Deze class dusdanig aanpassen, zodat de werkelijk tekst niet in het scripts staat, maar in templates in de projectmap.
De tekst kan op een soortgeljike manier worden opgeslagen in .txt bestanden met sepeciale tekens. zo kan str.replace()
worden toegepast om de data op de juiste plek in de tekst te substitueren.
"""


# todo: process schrijven voor wanneer verkeerde variabelen zijn ingevuld als parameters en er dus niets gebeurt
# todo: aanpassen in documentatie
class DocumentGenerator:

    def __init__(self, project: str, api_key: str, rapport_type: str, quarter: str, year: str, staging_file_name: str = None) -> None:
        self.sa = StoringsAnalyse(project, api_key, rapport_type, quarter, year, staging_file_name)
        self.newline = "\n"
        self.tab = "\t"

        self._default_export_file_name = f"TABLE_TEST_{self.sa.quarter}_{self.sa.year}_storingsanalyse_tekst.docx"
        self._default_export_file_name_appendix = f"{self.sa.quarter}_{self.sa.year}_storingsanalyse_bijlage.pdf"

        self._default_export_location = "documents/generated_documents"

    """
    Managing methods -- methods that fulfill some specific general task
    """
    def __aantal_per_maand(self, input_data: DataFrame, ntype: str) -> dict:
        meta_data_ntype = self.sa.metadata.return_ntype_meta_object(ntype=ntype)

        # Data voor in tekst ophalen
        ntype_per_maand = input_data['month_number'].value_counts()

        gemiddelde_per_maand = sum(ntype_per_maand) / len(ntype_per_maand)

        max_ntype_maand = self.sa.get_min_max_months(ntype_per_maand.to_dict(), min_max='max')
        min_ntype_maand = self.sa.get_min_max_months(ntype_per_maand.to_dict(), min_max='min')

        # todo: excluse year aanpassen zodat het dynamisch is en niet gewoon '2020'
        maanden = self.sa.metadata.get_month_list(exclude_year='2020')
        maandelijks_gemiddelde = self.sa.metadata.avg_monthly(dictionary=meta_data_ntype, exclude_keys=maanden)

        kwartaal_gemiddelde = self.sa.metadata.avg_quarterly(dictionary=meta_data_ntype)

        data_dict = {"ntype": ntype.lower(),
                     "totaal_aantal": len(input_data),
                     "ntype_count_per_maand": ntype_per_maand,
                     "gemiddelde_per_maand": gemiddelde_per_maand,  # just this quarter
                     "max_ntype_maand": max_ntype_maand,
                     "min_ntype_maand": min_ntype_maand,
                     "maandelijks_gemiddelde_start_project": maandelijks_gemiddelde,  # over the whole project
                     "kwartaal_gemiddelde": kwartaal_gemiddelde}

        return data_dict

    def build_table_docx(self, docx_object: docx.Document, headers: Tuple, row_data: List[Tuple[str, str]]) -> None:
        table = docx_object.add_table(1, cols=len(headers))
        table.style = docx_object.styles['Light Grid']
        # Adding headers
        row = table.rows[0].cells
        for i in range(len(headers)):
            row[i].text = headers[i]

        # Adding data from the input_dict
        for data in row_data:
            row = table.add_row().cells
            for i in range(len(data)):
                row[i].text = str(data[i])

        docx_object.save(self._default_export_file_name)

    def del_old_export(self):
        if self._default_export_file_name in os.listdir(os.curdir):
            print('Deleting old file')
            os.remove(self._default_export_file_name)
        else:
            pass

    """
    Chapter - Analyse
        Paragraph - Aantal meldingen
            Subsection - Aantal meldingen per maand
            Subsection - Aantal meldingen per subsysteem
        Paragraph - Aantal storingen
            Subsection - Aantal storingen per maand
            Subsection - Aantal storingen per subsysteem
    """
    def get_aantal_per_maand(self, ntype: str) -> dict:
        staging_data_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)
        return self.__aantal_per_maand(input_data=staging_data_ntype, ntype=ntype)

    def build_text_aantal_per_maand(self, input_data_dict: dict) -> str:
        text = f"""
        Om te kunnen bepalen of een trend waarneembaar is in het aantal meldingen per 
        maand, wordt als onderdeel van deze rapportage een grafiek toegevoegd. Zie 
        bijlage: “Aantal {input_data_dict["ntype"]} per maand”.

        Uit de grafiek valt het volgende te constateren:

        • Het totaal aantal {input_data_dict["ntype"]} in {self.sa.quarter} {self.sa.year} : {input_data_dict["totaal_aantal"]} 

        • Het gemiddelde aantal {input_data_dict["ntype"]} per maand : {input_data_dict["gemiddelde_per_maand"]}

        • Hoogste aantal {input_data_dict["ntype"]} in de maand{'en' if len(input_data_dict["max_ntype_maand"]) > 1 else ''} {', '.join(input_data_dict["max_ntype_maand"])}: {max(input_data_dict["ntype_count_per_maand"])}

        • Laagste aantal {input_data_dict["ntype"]} in de maand{'en' if len(input_data_dict["min_ntype_maand"]) > 1 else ''} {', '.join(input_data_dict["min_ntype_maand"])}: {min(input_data_dict["ntype_count_per_maand"])}

        • Het gemiddelde aantal {input_data_dict["ntype"]} per maand vanaf {self.sa.project_start_date}: {input_data_dict["maandelijks_gemiddelde_start_project"]}

        • Het gemiddelde aantal {input_data_dict["ntype"]} per kwartaal vanaf {self.sa.project_start_date}: {input_data_dict["kwartaal_gemiddelde"]}
        """
        return text

    def get_quarter_comparison(self, ntype: str) -> dict:
        """
        The subsections 'aantal ntype per maand' contain a comparison of the current quarter with the previous
        quarter, and a comparison of the current quarter with the selfme quarter of the previous year.
        :return:
        """
        staging_data_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)
        meta_data_ntype = self.sa.metadata.return_ntype_meta_object(ntype=ntype)

        # Aantal ntype in voorgaande q
        monthlist = self.sa.metadata.get_keys(dictionary=meta_data_ntype, containing_quarter=[self.sa.prev_quarter],
                                              containing_year=[self.sa.prev_year])
        ntype_gefilterd = self.sa.metadata.filter_dictionary_keys(dictionary=meta_data_ntype, keys=monthlist)
        totaal_ntype_voorgaand_kwartaal = self.sa.metadata.sum_values(ntype_gefilterd)

        # Aantal ntype in zelfde q voorgaand jaar
        monthlist = self.sa.metadata.get_keys(dictionary=meta_data_ntype, containing_quarter=[self.sa.quarter],
                                              containing_year=[self.sa.prev_year])
        ntype_gefilterd = self.sa.metadata.filter_dictionary_keys(dictionary=meta_data_ntype, keys=monthlist)
        totaal_ntype_zelfde_kwartaal = self.sa.metadata.sum_values(ntype_gefilterd)

        data_dict = {"ntype": ntype.lower(),
                     "totaal_aantal": len(staging_data_ntype),
                     "totaal_ntype_zelfde_kwartaal": totaal_ntype_zelfde_kwartaal,
                     "totaal_ntype_voorgaande_kwartaal": totaal_ntype_voorgaand_kwartaal}
        return data_dict

    def build_quarter_comparison(self, input_data_dict: dict) -> str:
        text = f"""
        In {self.sa.quarter} {self.sa.prev_year} waren in totaal {input_data_dict["totaal_ntype_zelfde_kwartaal"]} {input_data_dict["ntype"]} gemaakt. In {self.sa.quarter} {self.sa.year} zijn er {input_data_dict["totaal_aantal"] - input_data_dict["totaal_ntype_zelfde_kwartaal"]} {input_data_dict["ntype"]} 
        meer t.o.v. {self.sa.quarter} {self.sa.prev_year}. 

        In {self.sa.prev_quarter} {self.sa.year} waren in totaal {input_data_dict["totaal_ntype_voorgaande_kwartaal"]} {input_data_dict["ntype"]} gemaakt. In {self.sa.quarter} {self.sa.year} zijn er {input_data_dict["totaal_aantal"] - input_data_dict["totaal_ntype_voorgaande_kwartaal"]} {input_data_dict["ntype"]} 
        meer t.o.v. {self.sa.prev_quarter} {self.sa.year}. 
        """
        return text

    def get_aantal_per_subsysteem(self, ntype: str, threshold: int) -> dict:

        staging_data_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)

        # unieke types vastlegen
        # unique_types = list(staging_data_ntype.loc[:, 'sbs'].unique())

        sbs_count = staging_data_ntype.loc[:, 'sbs'].value_counts()

        # sbs nummers die verwerkt moeten worden
        sbs_to_process = [x for x in sbs_count.index if sbs_count.at[x] >= threshold]

        # lijst met rijen die verwerkt moeten worden
        rows_to_process = list()
        for sbs in sbs_to_process:
            ntype_per_sbs = sbs_count[sbs]
            percentage_ntype = round((ntype_per_sbs / sum(sbs_count)) * 100, 2)
            line = f"{self.sa.get_breakdown_description(sbs)}{self.tab}- {ntype_per_sbs} meldingen ({percentage_ntype}% van het totale aantal meldingen)"
            rows_to_process.append(line)

        # notification type
        ntypes = list(staging_data_ntype.loc[:, 'type melding (Storing/Incident/Preventief/Onterecht)'].unique())

        ntype_count = staging_data_ntype.loc[:, 'type melding (Storing/Incident/Preventief/Onterecht)'].value_counts()

        lines_to_process = list()
        for n in ntypes:
            line = f"{ntype_count[n]} {ntype.lower()} zijn gecategoriseerd als {n}."
            lines_to_process.append(line)

        data_dict = {"ntype": ntype.lower(),
                     "threshold": threshold,
                     "totaal_aantal": len(staging_data_ntype),
                     "sbs_count": sbs_count,
                     "sbs_to_process": sbs_to_process,
                     "rows_to_process": rows_to_process,
                     "lines_to_process": lines_to_process}
        return data_dict

    def build_text_aantal_per_subsysteem(self, input_data_dict: dict) -> str:
        text = f"""
        Er wordt en Pareto analyse gemaakt van het totaal aantal {input_data_dict['ntype']} per subsysteem. Deze is toegevoegd als bijlage. 

        Uit de pareto blijkt dat in {self.sa.quarter} {self.sa.prev_year} een totaal van {input_data_dict["totaal_aantal"]} {input_data_dict['ntype']} zijn gemeld, intern 
        dan wel extern. Voor het overzicht zijn de {input_data_dict['ntype']} bekeken met {input_data_dict["threshold"]} of meer 
        {input_data_dict['ntype']}. Dit is de top {len(input_data_dict["sbs_to_process"])} en heeft een totaal van {sum(input_data_dict["sbs_count"][input_data_dict["sbs_to_process"]])} {input_data_dict['ntype']} van de in totaal 
        {input_data_dict["totaal_aantal"]} (dit is {round((sum(input_data_dict["sbs_count"][input_data_dict["sbs_to_process"]]) / input_data_dict["totaal_aantal"]) * 100, 2)}% van het totaal). 
        Hieronder staan de deelinstallatie{'s' if len(input_data_dict["sbs_to_process"]) > 1 else ''}:


        {''.join((self.newline + '-' + self.tab + line + '' + self.newline for line in input_data_dict["rows_to_process"]))}


        De {input_data_dict["totaal_aantal"]} van {self.sa.quarter} {self.sa.year} zijn als volgt onder te verdelen:

        {''.join((self.newline + '-' + self.tab + ntype_line + '' + self.newline for ntype_line in input_data_dict["lines_to_process"]))}
        """
        return text

    """
    Chapter - Conclusie/Aanbeveling
        Paragraph - Algemeen
            Subsection - Probleem
            Subsection - Oorzaak
            Subsection - Oplossing
        Paragraph - Storingen per deelinstallatie
            Subsection - [subsectie voor elk uitgewerkte subsysteem. Title = subsystem_name]
    """
    def build_conclusie_algemeen_intro(self) -> str:
        staging_data_ntype = self.sa.return_ntype_staging_file_object(ntype='meldingen')
        sbs_count = staging_data_ntype.loc[:, 'sbs'].value_counts(dropna=False)

        text = f"""
        Er heeft een analyse van de storingen plaatsgevonden. Uit deze analyse is niet 
        naar voren gekomen dat verbeteren aan het onderhoudsplan en/of procedures en/of 
        hardware noodzakelijk zijn om het faalgedrag te verbeteren. 

        Alle meldingen moeten aan een asset / sub niveau van een DI worden gekoppeld. 
        Zodat altijd is te herleiden wat precies is gefaald. Aan alle meldingen is een DI 
        gekoppeld. Aan {max(sbs_count[sbs_count.index.isnull()].values) if len(sbs_count[sbs_count.index.isnull()]) != 0 else 0} werkorders zit geen sbs nummer gekoppeld. (zie besluit 5). 

        De meldingen zijn gekoppeld aan een probleem, oorzaak en oplossing. 

        Vanaf 1 september 2018 heeft een update plaats gevonden van het 
        onderhoudsmanagementsysteem. Bij deze update is het invullen van probleem, 
        oorzaak en oplossing toegevoegd in het systeem. Vanaf Q4 2018 zal dit ook 
        worden meegenomen in de analyse. In de volgende paragrafen staat de uitwerking 
        hiervan. Daarbij zie je het aantal van het huidige jaar, het totaal aantal en het 
        gemiddelde per Q vanaf Q4 2018.
        """
        return text

    def get_poo_table_data_md(self, poo_type: str) -> dict:
        """
        Retrieves/collects the data for the poo table in a way that is compatible wit the method that builds
        the markdown style table.
        :param poo_type:
        :return:
        """
        # todo: build a method to update the meta with the staging file poo data. this action gets easier when only
        #  needing to acces poo_from_meta
        poo_type_string = self.sa.metadata.return_poo_type_string(poo_type)

        poo_type_count = self.sa.meldingen[poo_type_string].value_counts(dropna=False).to_dict()

        meta_poo_type = self.sa.metadata.poo_data()[poo_type_string.split(' ')[0]]

        poo_type_avg_table = self.sa.metadata.poo_avg_table(poo_dictionary=meta_poo_type, poo_type=poo_type)

        poo_beschrijvingen = self.sa.metadata.contract_info()['POO_codes']

        data_dict = dict()
        for code in self.sa.metadata.return_poo_code_list(poo_type):
            code_counts = list()

            for quarter in meta_poo_type.keys():
                if code in meta_poo_type[quarter].keys():
                    code_counts.append(self.sa.metadata.sum_values(dictionary=meta_poo_type[quarter], keys=[code]))
                else:
                    code_counts.append(0)

            totaal = sum(code_counts)
            if code not in poo_type_count.keys():
                line = ''.join((self.newline + '|' + code + '|' + poo_beschrijvingen[code] + '|' + str(0) + '|' + str(totaal) + '|' + str(poo_type_avg_table[code]) + '|'))
                data_dict[code] = line
                continue

            line = ''.join((self.newline + '|' + code + '|' + poo_beschrijvingen[code] + '|' + str(poo_type_count[code]) + '|' + str(totaal + poo_type_count[code]) + '|' + str(poo_type_avg_table[code]) + '|'))
            data_dict[code] = line

        return data_dict

    @staticmethod
    def build_poo_table_md(input_data_dict: dict) -> str:
        """
        Build a poo data table in a markdown style.
        :param input_data_dict:
        :return:
        """
        text = """|Probleem|Beschrijving|Aantal|Totaal|Gemiddelde|
                  |--------|------------|------|------|----------|""" + ''.join((input_data_dict[code] for code in input_data_dict.keys()))
        return text

    def get_poo_table_data_v2(self, poo_type: str) -> dict:
        """
        Retrieves/collects the data for the poo table in a way that is compatible wit the method that builds
        the docx style table.
        :param poo_type:
        :return:
        """
        # todo: build a method to update the meta with the staging file poo data. this action gets easier when only
        #  needing to acces poo_from_meta
        poo_type_string = self.sa.metadata.return_poo_type_string(poo_type)

        poo_type_count = self.sa.meldingen[poo_type_string].value_counts(dropna=False).to_dict()

        meta_poo_type = self.sa.metadata.poo_data()[poo_type_string.split(' ')[0]]

        poo_type_avg_table = self.sa.metadata.poo_avg_table(poo_dictionary=meta_poo_type, poo_type=poo_type)

        poo_beschrijvingen = self.sa.metadata.contract_info()['POO_codes']

        data_dict = {"poo_type": poo_type_string.title(),
                     "rows": []}

        for code in self.sa.metadata.return_poo_code_list(poo_type):
            code_counts = list()

            for quarter in meta_poo_type.keys():
                if code in meta_poo_type[quarter].keys():
                    code_counts.append(self.sa.metadata.sum_values(dictionary=meta_poo_type[quarter], keys=[code]))
                else:
                    code_counts.append(0)

            aantal_count = poo_type_count[code] if code in poo_type_count.keys() else 0
            totaal = sum(code_counts) + aantal_count
            data_dict["rows"].append((code, poo_beschrijvingen[code], aantal_count, totaal, poo_type_avg_table[code]))

        return data_dict

    def build_poo_type_table(self, input_data: dict, docx_object: docx.Document) -> None:
        headers = (input_data["poo_type"], "Beschrijving", "Aantal", "Totaal", "Gemiddelde")
        self.build_table_docx(docx_object=docx_object, headers=headers, row_data=input_data["rows"])

    def get_aantal_per_subsysteem_per_maand(self, threshold: int, ntype: str = 'storingen') -> dict:
        """

        :param threshold:
        :param ntype: default value = 'storingen' because the standard analysis is preformed on the ntype 'storingen'
        :return:
        """
        staging_data_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)
        sf_data_groupby_sbs = staging_data_ntype.groupby('sbs')

        _group_data_to_print = dict()
        for group in sf_data_groupby_sbs.groups:
            group_data = sf_data_groupby_sbs.get_group(group)

            if len(group_data.index) < threshold:
                continue

            _group_data_to_print[group] = self.__aantal_per_maand(input_data=group_data, ntype=ntype)

        sorted_keys = sorted(_group_data_to_print, key=lambda item: _group_data_to_print[item]['totaal_aantal'], reverse=True)
        group_data_dict = {key: _group_data_to_print[key] for key in sorted_keys}
        return group_data_dict

    def build_aantal_per_subsysteem_per_maand(self, input_data_dict: dict) -> str:
        """

        :param input_data_dict: a dict of dicts for each subsystem a subdict
        :return:
        """
        text = """"""
        for sub_system in input_data_dict.keys():
            sub_system_data = input_data_dict[sub_system]
            sub_system_name = self.sa.get_breakdown_description(sbs_lbs=sub_system)
            text = text + '# ' + str(sub_system) + ' ' + sub_system_name + self.newline + self.build_text_aantal_per_maand(input_data_dict=sub_system_data) + self.newline

        return text

    """
    Chapter - Assets met de meeste meldingen
        Paragraph - Algemeen
        Paragraph - Uitwerking meldingen
        Paragraph - Conclusie
    """
    def get_asset_meeste_ntype_algemeen(self, threshold: int, ntype: str = 'meldingen') -> dict:
        staging_file_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)

        ntype_per_asset = staging_file_ntype['asset nummer'].value_counts()
        ntype_per_asset = ntype_per_asset.reset_index()
        ntype_per_asset.rename(columns={"asset nummer": "count", "index": "asset nummer"}, inplace=True)

        list_descriptions = {staging_file_ntype['asset nummer'][index]: staging_file_ntype['asset beschrijving'][index]
                             for index in range(staging_file_ntype.shape[0])}

        # asset beschrijving ophalen van de asset nummers om op een latere regel toe te voegen aan df ntype_per_asset
        asset_beschrijvingen = []
        for index, row in ntype_per_asset.iterrows():
            asset_num = row[0]
            if asset_num in list_descriptions.keys():
                asset_beschrijvingen.append(list_descriptions[asset_num])

        # print(f'\nasset beschrijving = {asset_beschrijvingen}\n')

        ntype_per_asset.at[:, 'asset beschrijving'] = asset_beschrijvingen

        # ophalen van de sbs nummers van de assets
        sbs_dict = dict()
        for asset_num in ntype_per_asset.loc[:, 'asset nummer'].to_dict().values():
            row = staging_file_ntype[staging_file_ntype.loc[:, 'asset nummer'] == asset_num]
            sbs_dict[asset_num] = row['sbs'].unique()[0]

        # bouwen van de lijst met regels data dide gepresenteerd moeten worden
        data_dict = {"threshold": threshold,
                     "lines": []}
        lines2handle = ntype_per_asset[ntype_per_asset['count'] >= threshold]
        for r in lines2handle.iterrows():
            row = r[1]
            line = ''.join((self.newline + '|' + str(self.sa.get_breakdown_description(sbs_dict[row[0]])) + '|' + str(row[-1]) + '|' + str(row[1]) + '|'))
            data_dict['lines'].append(line)

        return data_dict

    @staticmethod
    def build_asset_meeste_ntype_algemeen(input_dict: dict) -> str:
        text = """|Deelinstallatie|Asset|Aantal|
                  |---------------|-----|------|""" + ''.join((line for line in input_dict['lines']))
        return text

    def get_asset_meeste_ntype_algemeen_v2(self, threshold: int, ntype: str = 'meldingen') -> dict:
        staging_file_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)

        ntype_per_asset = staging_file_ntype['asset nummer'].value_counts()
        ntype_per_asset = ntype_per_asset.reset_index()
        ntype_per_asset.rename(columns={"asset nummer": "count", "index": "asset nummer"}, inplace=True)

        list_descriptions = {staging_file_ntype['asset nummer'][index]: staging_file_ntype['asset beschrijving'][index]
                             for index in range(staging_file_ntype.shape[0])}

        # asset beschrijving ophalen van de asset nummers om op een latere regel toe te voegen aan df ntype_per_asset
        asset_beschrijvingen = []
        for index, row in ntype_per_asset.iterrows():
            asset_num = row[0]
            if asset_num in list_descriptions.keys():
                asset_beschrijvingen.append(list_descriptions[asset_num])

        # print(f'\nasset beschrijving = {asset_beschrijvingen}\n')

        ntype_per_asset.at[:, 'asset beschrijving'] = asset_beschrijvingen

        # ophalen van de sbs nummers van de assets
        sbs_dict = dict()
        for asset_num in ntype_per_asset.loc[:, 'asset nummer'].to_dict().values():
            row = staging_file_ntype[staging_file_ntype.loc[:, 'asset nummer'] == asset_num]
            sbs_dict[asset_num] = row['sbs'].unique()[0]

        # bouwen van de lijst met regels data dide gepresenteerd moeten worden
        data_dict = {"threshold": threshold,
                     "rows": []}
        lines2handle = ntype_per_asset[ntype_per_asset['count'] >= threshold]
        for r in lines2handle.iterrows():
            row = r[1]
            data_dict["rows"].append((self.sa.get_breakdown_description(sbs_dict[row[0]]), row[-1], row[1]))

        return data_dict

    def build_asset_meeste_ntype_algemeen_v2(self, input_dict: dict, docx_object: docx.Document) -> None:
        headers = ("Deelinstallatie", "Asset", "Aantal")
        self.build_table_docx(docx_object=docx_object, headers=headers, row_data=input_dict['rows'])

    def get_asset_uitwerking_ntypes(self, threshold: int, ntype: str = 'meldingen') -> dict:
        staging_file_ntype = self.sa.return_ntype_staging_file_object(ntype=ntype)

        ntype_per_asset = staging_file_ntype['asset nummer'].value_counts(dropna=False)

        ntype_nan = ntype_per_asset.loc[np.nan] if np.nan in ntype_per_asset else 0

        ntype_per_asset = ntype_per_asset[(x is not np.nan for x in list(ntype_per_asset.index))]
        ntype_per_asset = ntype_per_asset.loc[ntype_per_asset >= threshold]

        data_dict = {"ntype": ntype,
                     "threshold": threshold,
                     "ntype_nan": ntype_nan,
                     "ntype_per_asset": ntype_per_asset}
        return data_dict

    def build_asset_uitwerking_ntypes(self, input_dict: dict):
        staging_file_ntype = self.sa.return_ntype_staging_file_object(ntype=input_dict['ntype'])
        text = f"""De assets met {input_dict['threshold']} of meer meldingen zijn hieronder uitgewerkt: 

        Bij de {input_dict['ntype_per_asset']} meldingen is geen asset gekoppeld aan de werkorder.
        """

        columns2present = ['werkorder', 'status', 'rapport datum', 'werkorder beschrijving', 'sbs',
                           'sbs omschrijving', 'locatie', 'locatie omschrijving', 'probleem code',
                           'beschrijving probleem', 'oorzaak code', 'beschrijving oorzaak',
                           'oplossing code', 'beschrijving oplossing', 'uitgevoerde werkzaamheden',
                           'type melding (Storing/Incident/Preventief/Onterecht)']

        for asset in list(input_dict['ntype_per_asset'].index):
            df = staging_file_ntype[staging_file_ntype["asset nummer"] == asset].copy().reset_index()
            line = f"""De {len(df)} meldingen van {df.loc[0, 'asset beschrijving']} worden hieronder gepresenteerd.
                       {self.newline}{df.loc[:, columns2present].to_html()}{self.newline}"""
            text = text + self.newline + line

        return text

    def build_asset_uitwerking_ntypes_v2(self, input_dict: dict, docx_paragraph_object, docx_object: docx.Document) -> None:
        """
        de v2 methods combineren het opstellen van de data met het schrijven naar het docx_object
        :param input_dict:
        :param docx_paragraph_object: type = docx.text.paragraph.Paragraph maar 'text' wordt niet herkend
        :param docx_object:
        :return:
        """
        staging_file_ntype = self.sa.return_ntype_staging_file_object(ntype=input_dict['ntype'])
        text = f"""{self.newline}De assets met {input_dict['threshold']} of meer meldingen zijn hieronder uitgewerkt: 

        Bij {sum(input_dict['ntype_per_asset'].to_dict().values())} meldingen is geen asset gekoppeld aan de werkorder.{self.newline}
        """

        docx_paragraph_object.add_run(text)

        columns2present = ['werkorder', 'rapport datum', 'type melding (Storing/Incident/Preventief/Onterecht)',
                           'werkorder beschrijving', 'probleem code', 'beschrijving probleem', 'oorzaak code',
                           'beschrijving oorzaak', 'oplossing code', 'beschrijving oplossing',
                           'uitgevoerde werkzaamheden']

        for asset in list(input_dict['ntype_per_asset'].index):
            df = staging_file_ntype[staging_file_ntype["asset nummer"] == asset].copy().reset_index()
            docx_object.add_heading('Asset: ' + str(df.loc[0, 'asset beschrijving']), level=3)

            # Onderstaande tekst komt niet op de plaats waar deze moet komen
            # todo: bovenstaande
            line = f"""De {len(df)} meldingen van {df.loc[0, 'asset beschrijving']} worden hieronder gepresenteerd.{self.newline}"""
            docx_paragraph_object.add_run(line)
            docx_object.save(self._default_export_file_name)

            all_table_data = [[(str(column_name), str(df.at[idx, column_name])) for column_name in columns2present] for idx in range(len(df))]
            for table_data in all_table_data:
                docx_object.add_heading('', level=4)  # Hack om verschillende kolommen toe te kunnen voegen
                self.build_table_docx(docx_object=docx_object, headers=('Kolom', 'Waarde'), row_data=table_data)
                # Above warning: Expected type 'Iterable[str]', got 'List[Tuple[str, str]]' instead. Can be ignored
                # docx_paragraph_object.add_run(f"""{self.newline}""")
            docx_object.save(self._default_export_file_name)

    @staticmethod
    def build_asset_conclusie(input_dict: dict) -> str:
        """

        :param input_dict: use datadict returned by get_asset_meeste_ntype_algemeen()
        :return:
        """
        text = f"""Als wordt gekeken naar de oorzaken van de meldingen van de {len(input_dict['lines'])} assets welke {input_dict['threshold']} of 
        meerdere meldingen hebben gehad, is bij somige assets repeterend en bij andere telkens verschillen.  

        Het falen van deze assets hoeft niet verder worden bekeken of worden onderzocht. 
        Dit omdat deze al reeds zijn behandeld bij de verschillende systemen. 
        """
        return text

    """
    Full document builders - methods that are responsible 
    """
    # todo: aanpassen in documentatie
    def build_full_document(self, path_to_folder: str = "", threshold: int = 3):
        # todo: aanpassen zodat de bold koppen level 3 koppen worden
        self.del_old_export()
        print('Creating file ' + self._default_export_file_name)

        doc = docx.Document()

        doc.add_heading("Analyse", level=1)
        doc.add_heading("Aantal meldingen", level=2)

        meldingen_per_maand = doc.add_paragraph("")
        meldingen_per_maand.add_run("Aantal meldingen per maand").bold = True
        meldingen_per_maand.add_run(self.build_text_aantal_per_maand(self.get_aantal_per_maand(ntype='meldingen')))
        meldingen_per_maand.add_run(self.build_quarter_comparison(self.get_quarter_comparison(ntype='meldingen')))
        doc.save(self._default_export_file_name)

        meldingen_per_subsysteem = doc.add_paragraph("")
        meldingen_per_subsysteem.add_run("Aantal meldingen per subsysteem").bold = True
        meldingen_per_subsysteem.add_run(self.build_text_aantal_per_subsysteem(self.get_aantal_per_subsysteem(ntype='meldingen', threshold=threshold)))
        doc.save(self._default_export_file_name)

        doc.add_heading("Aantal storingen", level=2)

        storingen_per_maand = doc.add_paragraph("")
        storingen_per_maand.add_run("Aantal storingen per maand").bold = True
        storingen_per_maand.add_run(self.build_text_aantal_per_maand(self.get_aantal_per_maand(ntype='storingen')))
        storingen_per_maand.add_run(self.build_quarter_comparison(self.get_quarter_comparison(ntype='storingen')))
        doc.save(self._default_export_file_name)

        storingen_per_subsysteem = doc.add_paragraph("")
        storingen_per_subsysteem.add_run("Aantal storingen per subsysteem").bold = True
        storingen_per_subsysteem.add_run(self.build_text_aantal_per_subsysteem(self.get_aantal_per_subsysteem(ntype='storingen', threshold=threshold)))
        doc.save(self._default_export_file_name)

        doc.add_heading("Conclusie/Aanbeveling", level=1)
        doc.add_heading("Algemeen", level=2)

        # conclusie_algemeen = doc.add_paragraph("")
        # conclusie_algemeen.add_run(self.build_conclusie_algemeen_intro())
        # conclusie_algemeen.add_run("\n")
        # conclusie_algemeen.add_run("Probleem").bold = True
        # conclusie_algemeen.add_run("\n")
        # # conclusie_algemeen.add_run(self.build_poo_table_md(self.get_poo_table_data_md('probleem')))
        #
        # conclusie_algemeen.add_run("\n")
        # conclusie_algemeen.add_run(self.newline)
        # doc.save(self._default_export_file_name)
        #
        # conclusie_algemeen.add_run("Oorzaak").bold = True
        # conclusie_algemeen.add_run("\n")
        # # conclusie_algemeen.add_run(self.build_poo_table_md(self.get_poo_table_data_md('oorzaak')))
        #
        # conclusie_algemeen.add_run("\n")
        # conclusie_algemeen.add_run("Oplossing").bold = True
        # conclusie_algemeen.add_run("\n")
        # # conclusie_algemeen.add_run(self.build_poo_table_md(self.get_poo_table_data_md('oplossing')))
        #
        # doc.save(self._default_export_file_name)

        conclusie_algemeen = doc.add_paragraph("")
        conclusie_algemeen.add_run(self.build_conclusie_algemeen_intro())
        conclusie_algemeen.add_run(self.newline)
        conclusie_algemeen.add_run(self.newline)
        doc.save(self._default_export_file_name)

        doc.add_heading("Probleem", level=2)
        self.build_poo_type_table(input_data=self.get_poo_table_data_v2(poo_type='probleem'),
                                  docx_object=doc)

        doc.add_heading("Oorzaak", level=2)
        self.build_poo_type_table(input_data=self.get_poo_table_data_v2(poo_type='oorzaak'),
                                  docx_object=doc)

        doc.add_heading("Oplossing", level=2)
        self.build_poo_type_table(input_data=self.get_poo_table_data_v2(poo_type='oplossing'),
                                  docx_object=doc)

        doc.add_heading("Storingen per deelinstallatie", level=2)

        storingen_per_subsysteem_per_maand = doc.add_paragraph("")
        storingen_per_subsysteem_per_maand.add_run("Uitwerking per deelinstallatie").bold = True
        storingen_per_subsysteem_per_maand.add_run(self.build_aantal_per_subsysteem_per_maand(self.get_aantal_per_subsysteem_per_maand(threshold=threshold)))
        doc.save(self._default_export_file_name)

        doc.add_heading("Assets met de meeste meldingen", level=1)
        doc.add_heading("Algemeen", level=2)

        self.build_asset_meeste_ntype_algemeen_v2(self.get_asset_meeste_ntype_algemeen_v2(threshold=threshold),
                                                  docx_object=doc)
        doc.save(self._default_export_file_name)

        doc.add_heading("Uitwerking meldingen", level=2)

        asset_uitwerking = doc.add_paragraph("")
        # asset_uitwerking.add_run(self.build_asset_uitwerking_ntypes(self.get_asset_uitwerking_ntypes(threshold=threshold)))

        self.build_asset_uitwerking_ntypes_v2(self.get_asset_uitwerking_ntypes(threshold=threshold),
                                              docx_paragraph_object=asset_uitwerking,
                                              docx_object=doc)

        doc.save(self._default_export_file_name)

        doc.add_heading("Conclusie", level=2)

        asset_conclusie = doc.add_paragraph("")
        # asset_conclusie.add_run("Conclusie").bold = True
        # conclusie_algemeen.add_run(self.newline)
        # conclusie_algemeen.add_run(self.newline)
        # doc.save(self._default_export_file_name)
        asset_conclusie.add_run(self.build_asset_conclusie(self.get_asset_meeste_ntype_algemeen(threshold=threshold)))
        doc.save(os.path.join(path_to_folder, self._default_export_file_name))

        print('Done. The text file is stored at ' + os.getcwd())

    def build_appendix(self, path_to_folder: str = "", threshold: int = 0):
        print('Creating file ' + self._default_export_file_name_appendix)
        title = 'title'
        #
        # Aantal meldingen per deelinstallatie
        #
        df = self.sa.return_ntype_staging_file_object(ntype='meldingen')
        time_range = [min(df['rapport datum']), max(df['rapport datum'])]
        available_categories = self.sa.metadata.contract_info()['aanwezige_deelinstallaties']

        categories, prepped_data = self.sa.prep(df, time_range, available_categories,
                                                time_key='rapport datum', category_key='sbs')

        readable_labels = [self.sa.prettify_time_label(label) for label in self.sa.last_seen_bin_names]
        self.sa.plot(input_data=prepped_data, plot_type='stacked',
                     category_labels=categories, bin_labels=readable_labels, title=title)

        summary_data = self.sa.prep_summary(df, time_range, available_categories, time_key='rapport datum',
                                            category_key='sbs')
        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        #
        # Aantal storingen per deelinstallatie
        #
        df = self.sa.return_ntype_staging_file_object(ntype='storingen')

        categories, prepped_data = self.sa.prep(df, time_range, available_categories,
                                                time_key='rapport datum', category_key='sbs')

        # needed to cover 'nan', else ValueError: shape mismatch: objects cannot be broadcast to a single shape
        readable_labels = [self.sa.prettify_time_label(label) for label in self.sa.last_seen_bin_names]
        self.sa.plot(input_data=prepped_data, plot_type='stacked',
                     category_labels=categories, bin_labels=readable_labels,
                     title=title)

        summary_data = self.sa.prep_summary(df, time_range, available_categories, time_key='rapport datum',
                                            category_key='sbs')
        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        #
        # Aantal onterechte meldingen per deelinstallatie
        #
        df = self.sa.return_ntype_staging_file_object(ntype='onterecht')

        categories, prepped_data = self.sa.prep(df, time_range, available_categories,
                                                time_key='rapport datum', category_key='sbs')

        readable_labels = [self.sa.prettify_time_label(label) for label in self.sa.last_seen_bin_names]

        self.sa.plot(input_data=prepped_data, plot_type='stacked',
                     category_labels=categories, bin_labels=readable_labels,
                     title=title)

        summary_data = self.sa.prep_summary(df, time_range, available_categories, time_key='rapport datum',
                                            category_key='sbs')
        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        #
        # Totaal aantal meldingen preventief per deelinstallatie
        #
        df = self.sa.return_ntype_staging_file_object(ntype='preventief')

        categories, prepped_data = self.sa.prep(df, time_range, available_categories,
                                                time_key='rapport datum', category_key='sbs')

        readable_labels = [self.sa.prettify_time_label(label) for label in self.sa.last_seen_bin_names]

        self.sa.plot(input_data=prepped_data, plot_type='stacked',
                     category_labels=categories, bin_labels=readable_labels,
                     title=title)

        summary_data = self.sa.prep_summary(df, time_range, available_categories, time_key='rapport datum',
                                            category_key='sbs')
        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        #
        # Aantal incidenten per deelinstallatie
        #
        df = self.sa.return_ntype_staging_file_object(ntype='incident')

        categories, prepped_data = self.sa.prep(df, time_range, available_categories,
                                                time_key='rapport datum', category_key='sbs')

        readable_labels = [self.sa.prettify_time_label(label) for label in self.sa.last_seen_bin_names]

        self.sa.plot(input_data=prepped_data, plot_type='stacked',
                     category_labels=categories, bin_labels=readable_labels,
                     title=title)

        summary_data = self.sa.prep_summary(df, time_range, available_categories, time_key='rapport datum',
                                            category_key='sbs')
        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        # todo: tijden dynamisch maken
        # Vergelijking voorgaande kwartaal met huidinge kwartaal
        #
        # Meldingen
        #
        categories, prepped_data = self.sa.prep(self.sa.metadata.unsaved_updated_meta['meldingen'],
                                                time_range=self.sa.get_time_range_v2(mode='pc'),
                                                available_categories=available_categories,
                                                time_key='rapport datum',
                                                category_key='sbs',
                                                bin_size='quarter')

        self.sa.plot(input_data=prepped_data,
                     plot_type='side-by-side',
                     category_labels=categories,
                     bin_labels=self.sa.last_seen_bin_names,
                     title=title)

        summary_data = self.sa.prep_summary(self.sa.metadata.unsaved_updated_meta['meldingen'],
                                            time_range=self.sa.get_time_range_v2(mode='pc'),
                                            available_categories=available_categories,
                                            bin_size='quarter')

        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        #
        # Storingen
        #
        categories, prepped_data = self.sa.prep(self.sa.metadata.unsaved_updated_meta['storingen'],
                                                time_range=self.sa.get_time_range_v2(mode='pc'),
                                                available_categories=available_categories,
                                                time_key='rapport datum',
                                                category_key='sbs',
                                                bin_size='quarter')

        self.sa.plot(input_data=prepped_data,
                     plot_type='side-by-side',
                     category_labels=categories,
                     bin_labels=self.sa.last_seen_bin_names,
                     title=title)

        summary_data = self.sa.prep_summary(self.sa.metadata.unsaved_updated_meta['storingen'],
                                            time_range=self.sa.get_time_range_v2(mode='pc'),
                                            available_categories=available_categories,
                                            bin_size='quarter')

        self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                             data=list(summary_data.values()),
                             title=title)

        #
        # Verdeling type meldingen per deelinstallatie
        #
        df = self.sa.return_ntype_staging_file_object(ntype='meldingen')
        df_groupby_sbs = df.groupby(['sbs'])

        # unieke types vastlegen
        unique_types = df.loc[:, 'type melding (Storing/Incident/Preventief/Onterecht)'].unique()

        # cols kan voor een sandbox tool variabel gemaakt worden.
        cols = ['type melding (Storing/Incident/Preventief/Onterecht)', 'month_number']

        sbs_count = df.loc[:, 'sbs'].value_counts()
        to_process = [x for x in sbs_count.index if sbs_count.at[x] >= threshold]
        for di_num in to_process:
            categories, prepped_data = self.sa.prep(df_groupby_sbs.get_group(di_num),
                                                    time_range=self.sa.get_time_range_v2(mode='pc'),
                                                    available_categories=unique_types,
                                                    time_key='rapport datum',
                                                    category_key='type melding (Storing/Incident/Preventief/Onterecht)')

            self.sa.plot(input_data=prepped_data,
                         plot_type='stacked',
                         category_labels=categories,
                         bin_labels=[self.sa.prettify_time_label(label) for label in self.sa.last_seen_bin_names],
                         title=title)

            summary_data = self.sa.prep_summary(df_groupby_sbs.get_group(di_num),
                                                time_range=self.sa.get_time_range_v2(mode='pc'),
                                                available_categories=unique_types,
                                                time_key='rapport datum',
                                                category_key='type melding (Storing/Incident/Preventief/Onterecht)')

            self.sa.plot_summary(x_labels=[self.sa.prettify_time_label(label) for label in summary_data.keys()],
                                 data=list(summary_data.values()),
                                 title=title)

        #
        # Exporting appendix
        #
        self.sa.export_graphs(filename=os.path.join(path_to_folder, self._default_export_file_name_appendix))
        print('Done. The appendix file is stored at ' + os.getcwd())


def main():
    """
    function to include the main functionality. If not defined like this and everything is imported using *. all
    variables defined in if __name__ == '__main__': will be imported too and clutters the executing code (maybe
    also introduce bugs).
    :return:
    """
    dg = DocumentGenerator(project="Coentunnel-tracé",
                           rapport_type="Kwartaalrapportage",
                           quarter="Q2",
                           year="2021",
                           api_key='bWF4YWRtaW46R21iQ1dlbkQyMDE5',
                           staging_file_name='validating_input_data.xlsx')
    dg.build_full_document()
    dg.build_appendix()


if __name__ == '__main__':
    os.chdir('..')
    main()
